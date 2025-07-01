import os
import re
from pathlib import Path
from docx import Document

RE_TOPIC = re.compile(r"^Topic Title:\s*(.+)$", re.IGNORECASE)
RE_SPEAKER = re.compile(r"^Speaker:\s*(.+)$", re.IGNORECASE)
RE_DATE = re.compile(r"^Date:\s*(.+)$", re.IGNORECASE)
SEPARATOR = "---"

RAW_DIR = Path("data/raw")
VALIDATED_DIR = Path("data/validated")
VALIDATED_DIR.mkdir(parents=True, exist_ok=True)

def load_paragraphs(path):
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def validate_structure(paragraphs):
    entries = []
    warnings = []

    current = {"topic": None, "speaker": None, "date": None, "body": []}

    for line in paragraphs + [SEPARATOR]:  # Final flush
        if line == SEPARATOR:
            if current["topic"] and current["body"]:
                entry = {
                    "topic": current["topic"],
                    "speaker": current["speaker"] or "unknown",
                    "date": current["date"] or "unknown",
                    "body": current["body"]
                }

                if not current["speaker"]:
                    warnings.append(f"⚠️ Missing speaker in topic: '{current['topic']}' → replaced with 'unknown'")
                if not current["date"]:
                    warnings.append(f"⚠️ Missing date in topic: '{current['topic']}' → replaced with 'unknown'")

                entries.append(entry)
            else:
                if not current["topic"]:
                    warnings.append("⚠️ Missing Topic Title in an entry → entry skipped.")
                elif not current["body"]:
                    warnings.append(f"⚠️ Missing body text for topic: '{current['topic']}' → entry skipped.")

            current = {"topic": None, "speaker": None, "date": None, "body": []}
            continue

        if match := RE_TOPIC.match(line):
            current["topic"] = match.group(1)
        elif match := RE_SPEAKER.match(line):
            current["speaker"] = match.group(1)
        elif match := RE_DATE.match(line):
            current["date"] = match.group(1)
        else:
            current["body"].append(line)

    return entries, warnings

def save_cleaned_doc(entries, output_path):
    doc = Document()
    for entry in entries:
        doc.add_paragraph(f"Topic Title: {entry['topic']}")
        doc.add_paragraph(f"Speaker: {entry['speaker']}")
        doc.add_paragraph(f"Date: {entry['date']}")
        for line in entry["body"]:
            doc.add_paragraph(line)
        doc.add_paragraph(SEPARATOR)
    doc.save(output_path)

def validate_all_docs():
    docx_files = list(RAW_DIR.glob("*.docx"))
    failed = []

    for file_path in docx_files:
        print(f"\n🔍 Validating: {file_path.name}")
        try:
            paragraphs = load_paragraphs(file_path)
            entries, warnings = validate_structure(paragraphs)

            if not entries:
                failed.append(file_path.name)
                print(f"❌ No valid entries found in {file_path.name} — file skipped.\n")
                continue

            output_path = VALIDATED_DIR / f"{file_path.stem}.validated.docx"
            save_cleaned_doc(entries, output_path)

            print(f"✅ {len(entries)} entries saved to {output_path.name}")
            for i, entry in enumerate(entries, 1):
                print(f"  Entry {i}: {entry['topic']} | {entry['speaker']} | {entry['date']}")

            if warnings:
                print("\n⚠️ Warnings:")
                for w in warnings:
                    print(" -", w)
            else:
                print("✅ No structural warnings.")
        except Exception as e:
            failed.append(file_path.name)
            print(f"❌ Error validating {file_path.name}: {e}")

    if failed:
        print("\n📛 The following files failed or had no valid entries:")
        for f in failed:
            print(" -", f)
    else:
        print("\n🎉 All files validated successfully!")

if __name__ == "__main__":
    validate_all_docs()


